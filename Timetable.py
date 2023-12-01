from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE
from docx.dml.color import ColorFormat
from docx.enum.dml import MSO_COLOR_TYPE
from docx.enum.text import WD_COLOR_INDEX, WD_COLOR, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml, OxmlElement
import random
import string


dayNames = {
  "0": "",
  "1": "Day 1",
  "2": "Day 2",
  "3": "Day 3",
  "4": "Day 4",
  "5": "Day 5",
  "6": "Day 6",
  "7": "Day 7",
  "8": "Day 8",
  "9": "Day 9",
  "10": "Day 10"
}
sessionNames = {
  "0": "",
  "1": "Care Group",
  "2": "Session 1",
  "3": "Recess",
  "4": "Session 2",
  "5": "Session 3",
  "6": "Lunch",
  "7": "Session 4"
}

def createBoldRun(paragraph, text):
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(8)
    run.bold = True
    return run


def createSmallRun(paragraph, text):
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(8)
    return run


def create(username, timetable):
    doc = Document()

    # Width & Margins
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Table Setup
    num_rows = 11
    num_cols = 8
    table = doc.add_table(rows=num_rows, cols=num_cols)

    # Create Table
    for i in range(num_rows):
        for j in range(num_cols):
            cell = table.cell(i, j)

            if j == 0:
                key = str(i)
                cell.text = dayNames[key]

            if i == 0:
                key = str(j)
                cell.text = sessionNames[key]

            if i == 0 or j == 0:
                shading_elm = parse_xml(r'<w:shd {} w:fill="9fc5e8"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            if i != 0 and j == 3 or i != 0 and j == 6:
                shading_elm = parse_xml(r'<w:shd {} w:fill="fff2cc"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)


        
            for loopNumber in dayNames:
                if loopNumber != "0":
                    if int(loopNumber) == i and i != 0:
                        
                        
                        for loopSecondNumber in sessionNames:
                            if loopSecondNumber != "0" and loopSecondNumber != "3" and loopSecondNumber != "6":
                                if int(loopSecondNumber) == j:
                                    
                                    classIndex = 0
                                    if int(loopSecondNumber) == 1:
                                        classIndex = 0
                                    elif int(loopSecondNumber) == 2:
                                        classIndex = 1
                                    elif int(loopSecondNumber) == 4:
                                        classIndex = 3
                                    elif int(loopSecondNumber) == 5:
                                        classIndex = 4
                                    elif int(loopSecondNumber) == 7:
                                        classIndex = 7

                                    paragraph = cell.paragraphs[0]
                                    bold_run = createBoldRun(paragraph, timetable[str(loopNumber)]['d']['Periods'][classIndex]['Classes'][0]['TimeTableClass'])
                                    paragraph.paragraph_format.space_before = Pt(2)
                                    paragraph.paragraph_format.space_after = Pt(0)

                                    paragraph = cell.add_paragraph()
                                    smaller_run = createSmallRun(paragraph, f"{timetable[str(loopNumber)]['d']['Periods'][classIndex]['Classes'][0]['TeacherName']}")
                                    paragraph.paragraph_format.space_before = Pt(0)
                                    paragraph.paragraph_format.space_after = Pt(0)

                                    paragraph = cell.add_paragraph()
                                    smaller_run = createSmallRun(paragraph, f"{timetable[str(loopNumber)]['d']['Periods'][classIndex]['Classes'][0]['Room']}")
                                    paragraph.paragraph_format.space_before = Pt(0)
    # Save File
    characters = string.ascii_letters + string.digits
    randomCode = ''.join(random.choice(characters) for _ in range(15))
    doc.save(f"./static/timetables/Timetable_{username}_{randomCode}.docx")
    return f"Timetable_{username}_{randomCode}.docx"